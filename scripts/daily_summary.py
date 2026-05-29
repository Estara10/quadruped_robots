"""Generate daily work summary as docx."""
import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_summary(date_str, content_md, output_dir=None):
    doc = Document()

    # Title
    title = doc.add_heading(f'ABS 项目日报 — {date_str}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # Content sections split by ## headers
    lines = content_md.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    # Footer
    doc.add_paragraph('')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')

    if output_dir is None:
        output_dir = os.path.expanduser('~/quadruped_robots/日报/')
    os.makedirs(output_dir, exist_ok=True)

    filename = f'日报_{date_str.replace("-", "")}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f'Saved: {filepath}')
    return filepath

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python daily_summary.py <date> <summary_text>')
        sys.exit(1)
    date_str = sys.argv[1]
    content = sys.argv[2]
    generate_summary(date_str, content)
