#!/usr/bin/env python3
"""Generate report: ABS paper reproduction status comparison."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Helper functions ──────────────────────────────

def set_run_font(run, cn_size=Pt(12), en_size=Pt(12), bold=False):
    """Set Chinese (宋体) and English (Times New Roman) font."""
    run.font.size = cn_size
    run.font.name = 'Times New Roman'
    run.bold = bold
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_paragraph_cn(text, size=Pt(12), bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     first_indent=True, space_after=Pt(0)):
    """Add a paragraph with Chinese formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.line_spacing = 1.5
    if first_indent:
        pf.first_line_indent = Pt(24)  # ~2 chars at 小四
    run = p.add_run(text)
    set_run_font(run, size, size, bold)
    return p

def add_heading_1(text):
    """一级标题: 居中, 小三, 加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, Pt(15), Pt(15), bold=True)
    return p

def add_heading_2(text):
    """二级标题: 四号, 加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, Pt(14), Pt(14), bold=True)
    return p

def add_heading_3(text):
    """三级标题: 小四, 加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, Pt(12), Pt(12), bold=True)
    return p

def add_table_row(table, cells_data, bold=False, header=False):
    """Add a row to table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.2
        pf.first_line_indent = Pt(0)
        run = p.add_run(str(text))
        set_run_font(run, Pt(10), Pt(10), bold=bold or header)

# ── Cover page ────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ABS 论文复现工作进展报告')
set_run_font(run, Pt(22), Pt(22), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Agile But Safe: Collision-Free High-Speed Quadruped Locomotion')
set_run_font(run, Pt(14), Pt(14))

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 年 6 月')
set_run_font(run, Pt(14), Pt(14))

doc.add_page_break()

# ── 一、工作概述 ──────────────────────────────────
add_heading_1('一、工作概述')

add_paragraph_cn(
    '本工作旨在复现 ABS（Agile But Safe）论文提出的双策略安全敏捷四足机器人运动控制框架。'
    '该框架发表于 RSS 2024，核心思想是通过敏捷策略与恢复策略的协同工作，'
    '使四足机器人在高速运动中具备自主避障能力。'
    '原论文基于 Unitree Go1 机器人、ROS1 Noetic 与 Isaac Gym 仿真环境实现。'
    '本工作在 Unitree Go2 机器人平台上，将系统整体迁移至 ROS2 Humble，'
    '并以 MuJoCo 物理引擎替代 Isaac Gym 作为仿真验证环境。'
)

add_paragraph_cn(
    '截至本报告撰写时（2026 年 6 月），已完成仿真环境中的 ABS 核心算法管线搭建与端到端验证，'
    '包括敏捷策略推理、RA 值网络评估、基于梯度下降的恢复策略 twist 优化、'
    '世界坐标系下的目标点导航与到达检测等核心功能。'
    'RA 和 recovery 策略的线速度已修正为机体系速度，与训练端和 ROS1 部署保持一致。'
    '机器人能够在平坦场景中导航至目标点并站定，'
    '在障碍物场景中遇到障碍时可触发 RA 评估和自动恢复行为。'
    '当前阶段正在针对直线性与避障行为进行仿真校准和结构化评估。'
)

# ── 二、论文系统概述 ──────────────────────────────
add_heading_1('二、论文系统架构概述')

add_heading_2('（一）训练阶段')

add_paragraph_cn(
    '论文的训练阶段基于 Isaac Gym 仿真平台，包含四个可训练模块。'
    '敏捷策略是目标导向的强化学习策略，接收 61 维观测并输出 12 维关节目标位置，'
    '通过 PPO 算法训练。RA 值网络是一个三层的全连接网络，结构为 19 输入、'
    '两个 64 维隐藏层、1 维 Tanh 输出，用于预测当前状态下的碰撞风险。'
    '恢复策略是 twist 跟踪策略，接收 49 维观测，输出关节目标，同样通过 PPO 训练。'
    '射线预测网络是 ResNet18 结构，将深度相机图像映射为 11 条稀疏射线距离，'
    '为策略和 RA 网络提供外部感知信息。'
)

add_heading_2('（二）部署阶段')

add_paragraph_cn(
    '论文的 ROS1 部署运行于 Orin NX 嵌入式计算平台，配合 ZED mini 深度相机。'
    'ZED 相机同时提供深度图像和基于视觉 SLAM 的世界坐标系定位。'
    '深度图像经 ResNet18 射线预测网络转换为 11 条对数空间射线距离。'
    '主控制循环以 50 Hz 频率运行，惯性测量单元数据和深度图像同步采集，'
    '策略推理以 12.5 Hz 的降采样频率执行。'
    'RA 值网络每推理步评估一次安全性，当 RA 值超过阈值时触发恢复策略。'
    '恢复 twist 通过梯度下降在 RA 模型上优化得到，优化目标为最小化碰撞风险同时约束偏离当前位置的程度。'
)

# ── 三、复现工作进展 ──────────────────────────────
add_heading_1('三、复现工作进展')

add_heading_2('（一）已完成模块')

# Table: completed modules
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0]
for i, text in enumerate(['模块', '论文方法', '本工作实现', '一致性']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, Pt(10), Pt(10), bold=True)

completed = [
    ['敏捷策略推理', '61 维观测 → 12 关节目标', 'TorchScript 加载, 61 维观测, 关节顺序 remap', '一致'],
    ['恢复策略推理', '49 维观测, twist 为 commands', 'TorchScript 加载, 内联替代 agile action', '一致'],
    ['RA 值网络', '19→64→64→1 Tanh, 阈值 -0.05', 'TorchScript 加载, 每 RL 步推理', '一致'],
    ['Recovery Twist 优化', '梯度下降 ×3, loss = lam·max(ra+2eps,0) + 0.02·pos_dev²', 'C++ torch::autograd, 3 迭代, lr=0.5, grad clip', '一致'],
    ['Recovery 触发', 'ra > -twist_eps (= -0.05)', 'ra > -0.05 进入, ra < -0.08 退出', '一致（加滞回）'],
    ['目标点导航', '世界坐标目标 → 机体坐标系位置指令', 'MuJoCo odometer + IMU 四元数 yaw', '一致（定位来源不同）'],
    ['到达检测', '训练 sigma_tight = 0.5m', 'dist < 0.5m 时 commands 清零', '一致'],
    ['Goal 重采样', '训练 episode reset 随机采样', '可配置开关, 默认首目标停止; 开启后到达 1.6s 站定重采样', '一致（行为可配）'],
    ['Contact 检测', '足力 > 1N (训练)', '足力 > 1N (MuJoCo touch sensor)', '一致'],
    ['射线感知 (仿真)', '几何射线对障碍物 actor', '2D 射线圆相交, geom 类型过滤', '功能等价'],
    ['Timer', '恒 0.5 (ROS1 部署)', '恒 0.5', '一致'],
]
for row_data in completed:
    add_table_row(table, row_data)

add_paragraph_cn('')

add_heading_2('（二）架构适配（ROS1→ROS2）')

add_paragraph_cn(
    'ROS1 到 ROS2 的迁移是项目面临的首要工程挑战。'
    '原论文的部署代码全部以 Python 编写为单文件脚本，运行于 ROS1 Noetic。'
    '本工作将其重构为 ROS2 Humble 下的 C++ 控制器插件，采用 ros2_control 框架的 FSM 状态机架构。'
    '推理引擎从 ONNX Runtime 更换为 LibTorch（TorchScript），以更好地与 C++ 代码集成。'
    '硬件接口层基于 unitree_sdk2 的 DDS 通信实现仿真器与控制器之间的数据交换。'
    '关节顺序方面，训练框架 Isaac Gym 以字母序导出关节名称，导致策略期望 FL 优先的关节排列，'
    '而本工作控制器按照 MuJoCo 模型的 FR 优先顺序组织数据，'
    '因此需要在观测构造和动作输出时进行关节顺序重映射。'
)

add_heading_2('（三）频率适配')

add_paragraph_cn(
    '原论文 ROS1 部署的主循环以 50 Hz 运行，策略推理降采样至 12.5 Hz，'
    '每推理步间隔 80 ms。本工作控制器以 500 Hz 运行，RL 推理降采样至 125 Hz，'
    '每推理步间隔仅 8 ms。频率差异导致恢复策略的行为显著不同：'
    'ROS1 中恢复策略持续 3 到 4 步约为 240 ms 到 320 ms，机器人有足够时间产生位移响应；'
    '本工作中若不处理，恢复策略仅持续 2 到 3 步即退出，总时间不足 25 ms，'
    '机器人尚未开始移动即被切换回敏捷策略，形成前冲后退的振荡。'
    '为此引入了频率适配的恢复保持机制，进入恢复后至少保持 30 个 RL 步约 240 ms，'
    '与 ROS1 的有效恢复时长对齐。保持期内复用梯度下降优化得到的同一个 twist，'
    '不再重新计算。'
)

add_heading_2('（四）未完成模块')

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0]
for i, text in enumerate(['模块', '论文方法', '未完成原因']):
    p = hdr2.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, Pt(10), Pt(10), bold=True)

pending = [
    ['射线预测网络 (ResNet18)', '深度相机 → ResNet18 → 11 射线', 'MuJoCo 渲染深度图与真实相机存在分布差异, 直接推理效果不佳, 需域适应或仿真重训练'],
    ['实机 Go2 部署', 'Ubuntu 20.04 + ROS1 + Orin NX + ZED mini', '硬件环境尚未就绪, 需逐步迁移至 Go2 + ROS2 + 待定计算平台'],
    ['遥控器紧急停止', '订阅 Unitree 无线遥控器 keydata', '仿真场景无物理遥控器, 实机部署时接入'],
    ['温度监控', '电机温度实时监测', '仿真中电机温度传感器无实际数据'],
    ['LED RA 反馈', 'LED 颜色指示 RA 值', '非功能性需求, 部署时添加'],
]
for row_data in pending:
    add_table_row(table2, row_data)

add_paragraph_cn('')

# ── 四、实验结果 ──────────────────────────────────
add_heading_1('四、实验结果')

add_heading_2('（一）平坦地面目标导航')

add_paragraph_cn(
    '在无任何障碍物的平坦地面场景中，机器人从原点出发，目标点设定为世界坐标 7 米正前方。'
    '测试中机器人沿直线前进，约 4 秒后到达目标点 0.5 米范围内，触发到达检测后站定。'
    '默认配置下到达首目标后停止，不再自动重采样新目标。'
    '整个过程中 RA 值稳定在 -0.7 以下，未出现误触发恢复的情况。'
    '验证了基于 odometer 的世界坐标系目标导航、机体系速度 RA 推理和到达检测功能正确。'
)

add_heading_2('（二）恢复策略梯度下降优化')

add_paragraph_cn(
    '在场景中存在障碍物时，RA 值上升超过 -0.05 的进入阈值，触发恢复策略。'
    '梯度下降优化从当前机器人速度初始化 twist，经过 3 次迭代后收敛。'
    '观察结果显示，vx 分量通常被推向负方向（后退），这是 RA 模型认为后退更安全所导致的结果；'
    'vy 分量被推向某个侧向极值（通常为 ±0.3 m/s）；wz 分量随 vy 方向相应调整。'
    '恢复保持机制使机器人在约 240 ms 内持续执行恢复策略动作，在此过程中 RA 值逐步下降。'
    '保持期结束后若 RA 低于退出阈值 -0.08，系统自动切回敏捷策略。'
)

add_heading_2('（三）目标重采样与连续导航')

add_paragraph_cn(
    '目标重采样功能默认为关闭状态（resample_goal_on_arrival: false），'
    '此时机器人到达第一个目标后持续输出站立指令，适合单次到达验证和调试。'
    '当配置开启时，机器人在到达目标点并站定后，自动在当前位置的前方 1.5 到 7.5 米、'
    '侧向 -2.0 到 2.0 米范围内随机采样下一个目标点。该采样范围与原论文训练时的目标分布一致。'
    '多目标点的连续导航测试表明，机器人能够依次完成目标切换并保持正常的避障行为。'
)

add_heading_2('（四）射线感知 geom 类型过滤')

add_paragraph_cn(
    'MuJoCo 场景中共有 162 个 geom 元素，包括地面平面、高度场地形、机器人自身的碰撞与视觉网格，'
    '以及显式放置的障碍物。原过滤器仅按 geom 组和动态质量排除机器人部件与地面平面，'
    '但无法区分障碍物 geom 与场景渲染产生的大量未命名视觉网格 geom，'
    '导致在平坦场景中也出现误检测。增加了按 geom 类型过滤的机制后，'
    '平面、高度场和网格类型的 geom 被排除，仅保留盒体、圆柱体、球体等显式障碍物类型。'
    '该过滤方式与原论文训练代码中仅对显式加载的障碍物 actor 进行射线查询的做法功能等价。'
)

# ── 五、分析与总结 ────────────────────────────────
add_heading_1('五、分析与总结')

add_heading_2('（一）与论文的一致性评价')

add_paragraph_cn(
    '在算法层面，本工作的实现与原论文保持了高度一致。'
    '四项核心推理模块——敏捷策略、恢复策略、RA 值网络和恢复 twist 优化——'
    '均严格按照论文 ROS1 部署代码中的方法实现。'
    'RA 值网络的触发阈值、恢复策略的内联替代方式、目标点导航的机体坐标系变换、'
    '距离缩放因子等细节均与原论文一致。'
    '主要差异来自 ROS1 到 ROS2 的工程迁移和仿真环境从 Isaac Gym 到 MuJoCo 的切换，'
    '这些差异属于必要的基础设施适配，不改变算法本身的逻辑。'
)

add_heading_2('（二）频率问题的处理')

add_paragraph_cn(
    'ROS2 控制器与原论文在运行频率上的差异是影响恢复行为的关键因素。'
    '本工作采用了保持步数机制来匹配原论文的有效恢复时长，这是频率差异下保持行为一致性的必要手段。'
    '该机制不影响论文的梯度下降优化逻辑，仅对切换时机做了时序层面的适配。'
)

add_heading_2('（三）机体系速度修正')

add_paragraph_cn(
    '在行为校准过程中发现 RA 值网络和恢复策略的线速度输入存在坐标系问题。'
    '训练端和 ROS1 部署均使用机体系速度（ZED 里程计经 world-to-body 旋转得到），'
    '而原始 ROS2 实现将 estimator 的世界系速度直接传给 RA 和 recovery 策略推理。'
    '这导致机器人在有 yaw 旋转时 RA 判断和 twist 优化失真，引起非必要偏航和避障行为异常。'
    '修正后优先使用 MuJoCo odometer 的世界系速度，再通过 IMU 四元数旋至机体系，'
    '与训练端 quat_rotate_inverse() 完全对齐。同时修正了状态估计器中 KDL 腿链顺序，'
    '使其与控制器 FR, FL, RR, RL 关节顺序一致，确保后备估计路径正确。'
)

add_heading_2('（四）后续工作方向')

add_paragraph_cn(
    '射线预测网络是连接仿真与实机的关键模块。当前在 MuJoCo 渲染深度图上直接使用 ResNet18 的效果'
    '不理想，原因是训练数据来自 ZED 真实相机，与渲染深度图之间存在分布差异。'
    '后续可以用 MuJoCo 仿真渲染的深度图重新训练或微调该模型，使其适应仿真环境，'
    '为最终部署到实机做好感知层面的准备。'
    '实机部署还需要完成硬件接口联调、深度相机集成、定位系统适配等工作。'
)

# ── Save ──────────────────────────────────────────
output_path = os.path.expanduser('~/quadruped_robots/ABS复现进展报告.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
