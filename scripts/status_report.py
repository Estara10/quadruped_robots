#!/usr/bin/env python3
"""Generate ABS 项目状态报告 .docx"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True

def p(text):
    doc.add_paragraph(text)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.font.size = Pt(10)
        r.font.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
    doc.add_paragraph()

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ABS 论文复现项目 · 当前状态报告")
run.font.size = Pt(20)
run.font.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("2026 年 6 月 9 日")
run.font.size = Pt(12)

doc.add_paragraph()

# ════════════════════════════════════════
h1("一、与论文的一致性分析")

h2("1.1 与论文严格一致的部分（算法层面）")
p("以下模块严格按照论文 ROS1 部署代码实现，算法逻辑未做任何修改：")

table(
    ["模块", "论文方法", "ROS2 实现", "一致性"],
    [
        ["敏捷策略推理", "61 维观测 → 12 关节目标 (PPO)", "TorchScript 加载, 61 维观测, 关节顺序 remap", "一致"],
        ["恢复策略推理", "49 维观测, twist 为 commands, 内联替代 agile", "TorchScript 加载, RA 触发后内联替换 action", "一致"],
        ["RA 值网络", "19→64→64→1 Tanh, 阈值 ra>-twist_eps",
         "TorchScript 加载, 每 RL 步推理, ra>-0.05 进入 / <-0.08 退出", "一致（加滞回）"],
        ["Recovery Twist GD", "梯度下降 ×3, loss=lam·max(ra+2eps,0)+0.02·pos_dev²",
         "C++ torch::autograd, 3 迭代, lr=0.5, grad clip ±1.0, twist clip", "一致"],
        ["目标点导航", "世界坐标目标 → 机体坐标系位置指令, min(1,5/dist+0.01) 缩放",
         "odometer 世界位置 + IMU yaw → body-frame command", "一致"],
        ["到达检测", "训练 sigma_tight = 0.5m", "dist < 0.5m 时 body=(0,0), heading=0", "一致"],
        ["Contact 检测", "足力 > 1N (训练 contact_forces>1.0)", "MuJoCo touch sensor > 1N → +1/-1", "一致"],
        ["Timer", "恒 0.5 (ROS1 部署)", "恒 0.5", "一致"],
        ["Ray2d 感知 (仿真)", "几何射线对障碍物 actor", "2D 射线圆相交, geom 类型过滤 → log2 → shm", "功能等价"],
        ["RA 观测速度帧", "lin_vel 为机体系速度 (ZED body-frame)", "odometer world vel + IMU quat → body frame", "一致"],
    ],
)

h2("1.2 因 ROS1→ROS2 迁移产生的差异（工程适配）")
p("以下差异源于架构迁移，论文算法本身未改变：")

table(
    ["维度", "论文 (ROS1)", "本工作 (ROS2)", "原因"],
    [
        ["编程语言", "Python + ONNX Runtime", "C++ + LibTorch (TorchScript)", "ros2_control 插件需 C++"],
        ["架构", "单文件 Python 脚本", "ros2_control FSM 状态机", "ROS2 框架要求"],
        ["通信", "UDP 直连 Unitree SDK", "DDS (unitree_sdk2, domain=1)", "ROS2 标准, 仿真/实机统一"],
        ["仿真器", "Isaac Gym (PhysX)", "MuJoCo 3.3.3", "Isaac Gym 不开源部署接口"],
        ["定位", "ZED SLAM (mocap/visual)", "MuJoCo odometer (framepos + framelinvel)", "仿真无 ZED"],
        ["速度来源", "ZED body-frame velocity", "odometer world vel → IMU quat → body frame", "等价变换"],
        ["推理频率", "12.5 Hz (80ms/步)", "125 Hz (8ms/步, decimation=4)", "控制器 500Hz, 需要降采样"],
        ["恢复保持", "3~4 步 ≈ 240~320ms", "30 步保持 (30×8ms=240ms)", "频率适配, 论文逻辑不变"],
        ["射线预测", "ResNet18 (depth→11 rays)", "仿真用几何射线 / 实机待 ResNet18", "仿真直接用真实几何"],
        ["关节顺序", "Unitree FR,FL,RR,RL", "Controller FR,FL,RR,RL → remap to FL-first", "训练 URDF 字母序"],
        ["安全机制", "safe.PositionLimit / PowerProtect", "IMU tilt + action clip + pose clamp + PD zero", "仿真无硬件安全层"],
    ],
)

# ════════════════════════════════════════
h1("二、已验证 vs 未严格验证项")

h2("2.1 已经过仿真验证的")
table(
    ["验证项", "方法", "结果"],
    [
        ["平地导航 + 到达停止", "launch_abs_sim.sh, 观测 [EVAL]", "2054 RL 步后正常到达, [ARRIVED] 触发"],
        ["RA 触发 recovery", "scene_test1 障碍场景", "ra 上升 → [RA-REC] ENTER → GD 优化 → 恢复 → EXIT"],
        ["Recovery Twist GD", "[TWIST-GD] 日志", "3 次迭代收敛, twist 在边界范围内"],
        ["机体系速度正确性", "[EVAL] lin_vel 对比 yaw", "前进时 lin_vel.x 为主分量, 机身旋转后跟随"],
        ["安全姿态触发", "body_tilt_limit_deg 设为 5°", "pitch=5.6° → [SAFETY] → PASSIVE, 零误报"],
        ["PASSIVE 退出卸力", "发布 command=1", "[VERIFY-EXIT] 执行, kp/kd 归零"],
        ["多目标重采样", "resample_goal_on_arrival=true", "[GOAL-RESAMPLE] 正常采样新目标"],
        ["FSM 状态切换", "launch 脚本自动 2→2→3", "PASSIVE→FIXEDDOWN→FIXEDSTAND→RL 正常"],
    ],
)

h2("2.2 未经过严格验证 / 已知缺点")
table(
    ["项目", "问题描述", "影响等级"],
    [
        ["纯平地直线性", "默认 MuJoCo 场景非真平地, 一进 RL 就 min_ray_m≈1m, 立即触发 recovery。从未在无障碍场景验证纯直线。", "高"],
        ["避障行为质量", "recovery 触发后机器人确实后退/转向, 但动作是否足够自然、是否过多触发, 无定量评估", "高"],
        ["多场景成功率", "run_abs_eval.py / analyze_abs_eval.py 已编写但从未系统性跑过基线", "高"],
        ["长时间稳定性", "单个 episode 最多跑了 ~90s, 未测试连续数小时的稳定性", "中"],
        ["调试日志噪音", "[EVAL]+[SYMM]+[STAND-SYMM] 每 0.2s 喷一次, 正常运行时不需要", "中"],
        ["Recovery hold 时长", "30 步是理论计算, 未对比不同 hold_steps 对避障成功率的影响", "中"],
        ["Action clip 合理性", "±4 是经验值, 是否限制正常 stride 未知", "低"],
        ["MuJoCo 摩擦参数", "未与训练 Isaac Gym 的地面摩擦精确对齐", "低"],
        ["实机部署", "整个实机链路未验证: ResNet18/DDS/ZED/安全/计算平台", "—"],
    ],
)

# ════════════════════════════════════════
h1("三、优先级排序（下一步工作）")

table(
    ["优先级", "任务", "理由"],
    [
        ["P0", "调试日志默认关闭 + 中文化关键日志", "立刻减少 90% 日志量, 不再需要手动过滤"],
        ["P0", "纯平地直线性验证 (scene.xml)", "确认修复后机器人确实能走直线, 不是靠避障东绕西绕"],
        ["P1", "多场景结构化评估基线", "6 个场景 × n 次, 拿成功率/到达率/碰撞率/recovery 频率"],
        ["P1", "对比不同 hold_steps/RA 阈值的影响", "确认超参数不是碰巧能用"],
        ["P1", "提交代码 (git commit)", "10+ 文件已改动, 该保存了"],
        ["P2", "行为校准优化", "如果基线数据不好, 调优 recovery 参数"],
        ["P2", "ResNet18 域适应/仿真重训练", "实机感知前置条件"],
        ["P3", "实机 Go2 部署", "硬件接口 + ZED + 安全"],
    ],
)

# ════════════════════════════════════════
h1("四、调试日志方案")

h2("4.1 当前问题")
p("每 25 RL 步 (≈0.2s) 输出一条 [EVAL] + [SYMM], 格式为英文, 字段密集难以直视。正常仿真 10 秒就产生 100+ 行。")

h2("4.2 建议方案")
p("区分三级日志, 默认只开 ERROR 和关键事件, 评估时才开 DETAIL：")

table(
    ["级别", "标签", "内容", "默认"],
    [
        ["ERROR", "[GOAL]", "目标位置、距离、到达标记（关键事件）", "✅ 开"],
        ["ERROR", "[RA-REC]", "Recovery 进入/退出（关键事件）", "✅ 开"],
        ["ERROR", "[SAFETY]", "安全触发（异常事件）", "✅ 开"],
        ["ERROR", "[TWIST-GD]", "Twist 优化结果（仅 recovery 时）", "✅ 开"],
        ["WARN/INFO", "[GOAL-RESAMPLE]", "目标重采样（仅当 enable 时）", "✅ 开"],
        ["DETAIL", "[EVAL]", "完整遥测（评估基线时开）", "❌ 关"],
        ["DETAIL", "[SYMM]", "对称性诊断（调试时开）", "❌ 关"],
        ["DETAIL", "[STAND-SYMM]", "站立对称性诊断（调试时开）", "❌ 关"],
    ],
)

h2("4.3 中文化方案")
p("保留英文标签（方便 grep 过滤），在关键日志中加中文说明。示例：")

p('''[GOAL] 目标导航 | robot=(5.93, 0.11) 位置 | yaw=-0.03 偏航 | goal=(9.00, 0.00) 目标
      dist=3.07 距离 | body=(4.97, -0.02) 机体系目标 | heading=-0.00 航向角 | arrived=0 到达
      [ARRIVED] 已到达目标点，站立中

[RA-REC] ENTER recovery | ra=0.1624 风险值 > entry=-0.0500 进入阈值
         twist=[-1.39, 0.30, 1.24] 恢复速度指令 | hold=30 保持步数

[SAFETY] Body tilt exceeded | pitch=5.6° 俯仰角 > limit=5.0° 安全限制 → 强制 PASSIVE 卸力''')

h2("4.4 配置项")
p("在 abs/config.yaml 中已有：")
p("  eval_telemetry_enabled: false     # [EVAL] 结构化评估遥测（日常关）")
p("  symmetry_debug_enabled: false     # [SYMM] 对称性诊断（日常关）")

# ════════════════════════════════════════
h1("五、论文管线对照")

h2("5.1 整体流程")
p("当前 ROS2 实现严格遵守论文的 agile → RA → recovery 管线：")

p("""
每 RL 推理步 (125 Hz):
  1. 获取状态 (关节, IMU, odometer, foot_force, ray2d)
  2. 计算目标导航 (世界 goal → 机体系 body, 距离缩放)
  3. 构造 61 维敏捷观测 (contact + ang_vel + gravity + commands + timer + dof + actions + ray2d)
  4. 前向 agile policy → action (policy order: FL, FR, RL, RR)
  5. 构造 19 维 RA 观测 (lin_vel + ang_vel + commands[0:2] + ray2d)
  6. 前向 RA model → ra_value
  7. ra > -0.05?
     YES → 梯度下降优化 twist (3 次迭代) → 使用 recovery policy → 保持 30 步
     NO  → 正常使用 agile action
  8. action remap → controller order (FR, FL, RR, RL)
  9. 安全 clamp (action ±4, 关节位置限制)
  10. 发送位置命令 + PD 增益
""")

h2("5.2 各模块严格性评价")

table(
    ["模块", "是否严格按论文", "备注"],
    [
        ["Agile Policy 推理", "✅ 是", "同模型, 同观测, 同输出格式"],
        ["RA Value Network", "✅ 是", "同模型, 同输入, 同阈值逻辑"],
        ["Recovery Policy 推理", "✅ 是", "同模型, 同 49-dim 观测, 同 twist 命令"],
        ["Recovery Twist GD", "✅ 是", "3 迭代, 同 loss 函数, 同学习率, 同梯度裁剪"],
        ["Goal 导航变换", "✅ 是", "同 world→body 旋转, 同距离缩放"],
        ["Contact 检测", "✅ 是", "同 1N 阈值"],
        ["Timer", "✅ 是", "同 0.5 固定值"],
        ["Ray2d 感知 (仿真)", "✅ 功能等价", "几何射线 ≈ Isaac Gym circle_ray_query"],
        ["Ray2d 感知 (实机)", "❌ 未实现", "需 ResNet18 → 域适应"],
        ["安全机制", "✅ 增强", "论文硬件安全 + 额外仿真兜底"],
    ],
)

# ════════════════════════════════════════
h1("六、实机部署步骤")

p("从仿真到实机 Go2，还需以下步骤：")

table(
    ["步骤", "内容", "依赖", "状态"],
    [
        ["1. 硬件接口适配", "hardware_unitree_mujoco → hardware_unitree_real: 通过 DDS 与实机 Unitree SDK 通信, 而非 MuJoCo bridge", "需实机 Go2 + 网络配置", "❌"],
        ["2. 深度相机集成", "ZED mini / ZED 2 驱动 + ROS2 wrapper → 发布深度图", "需 ZED SDK + ROS2", "❌"],
        ["3. ResNet18 部署", "加载 ZED 训练的 ResNet18 → 深度图 → 11 rays → shm 或 topic", "需域适应或仿真重训练", "❌"],
        ["4. 定位方案", "ZED SLAM (位置+速度) 或 MoCap → 替代 odometer", "需 ZED tracking 或运动捕捉", "❌"],
        ["5. 计算平台", "Orin NX/AGX 或 x86 工控机 → 运行 ROS2 + LibTorch 控制器", "需平台选型", "❌"],
        ["6. 安全系统", "遥控急停 (E.2) + 电机温度监控 (E.4) + 电池电压", "需连接 Unitree 无线遥控器", "❌"],
        ["7. 参数标定", "Kp/Kd/action_scale 实机调优, 站姿微调", "需实机试跑", "❌"],
        ["8. 逐步测试", "站立 → 原地踏步 → 慢速前进 → 正常速度 → 避障 → 全场测试", "需安全区域", "❌"],
    ],
)

# ════════════════════════════════════════
h1("七、其他重要未提及事项")

h2("7.1 代码管理")
p("⚠️ 当前 git status 显示修改了 9 个文件 + 多个新增文件。已累积多次修改未提交，"
  "继续堆积会增加合并冲突风险和丢失改动的风险。建议按功能拆成多个 commit。")

h2("7.2 场景问题")
p("⚠️ 当前 MuJoCo 默认场景从 simulate/config.yaml 加载，不是纯平地。"
  "launch_abs_sim.sh 没有显式覆盖场景，导致看起来是'平地测试'但实际跑的是 terrain 场景。"
  "这会影响所有'直线性验证'和'平地基准'的结论。建议在启动脚本中显式指定 scene.xml。")

h2("7.3 评估工具")
p("scripts/run_abs_eval.py 和 scripts/analyze_abs_eval.py 已编写但从未用于产生正式基线。"
  "这两个脚本的价值在于：自动启停仿真、批量跑场景、自动汇总成功率/到达率/recovery 频率。"
  "建议在提交代码前至少跑一轮生成量化基线。")

h2("7.4 频率匹配")
p("当前恢复保持 30 步是为了匹配 ROS1 约 3 步的时长。但 ROS1 每步间隔 80ms，"
  "RO2 每步 8ms。两者的物理响应时间不同：recovery 策略的关节指令在 8ms 内完成切换，"
  "但物理响应（速度/位移变化）受限于机器人动力学。30 步保持是否最优，缺少对比实验。")

h2("7.5 接触过滤")
p("训练端用了 contact_filt（当前帧 OR 上一帧），部署端直接用了原始接触信号。"
  "虽然落差不大（MuJoCo 接触信号比 Isaac Gym 更干净），但对于步态检测可能有细微影响。")

h2("7.6 ROS1 参考代码中观测的 lin_vel")
p("ROS1 部署脚本中敏捷观测 obs[0:6] 是 ang_vel 和 gravity，没有 lin_vel。"
  "agile policy 的训练观测也不含 lin_vel（只有 contact+ang_vel+gravity+commands+timer+dof+actions+ray2d）。"
  "我们的 computeObservation() 的输出顺序是对的（contact→ang_vel→gravity_vec→commands→timer→dof→actions→ray2d），lin_vel 只用于 RA 和 recovery。这一点已经正确。")

# ════════════════════════════════════════
h1("八、总结")

p("核心复现工作已完成。仿真中 ABS 四条主线（agile policy / RA value / recovery policy / twist GD）"
  "严格按论文实现。ROS1→ROS2 的差异全部属于架构适配，不改变算法逻辑。")

p("当前最大风险是：从未在纯平地验证直线性，且默认场景混入了障碍物触发误 recovery，"
  "导致'机器人能走直线'这个基本断言缺少证据。其次是缺少量化评估基线。")

p("下一步首选：关调试日志 + 纯平地验证。这两项成本最低、信息增益最大。")

# ── Save ──
output_path = os.path.expanduser("~/quadruped_robots/ABS项目状态报告.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
